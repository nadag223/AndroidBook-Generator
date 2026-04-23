import os
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt


# ─────────────────────────────
# לוגיקה
# ─────────────────────────────
EXCLUDED_DIRS     = {"build", "generated", ".gradle", ".idea", "test", "tests", "androidTest"}
EXCLUDED_KEYWORDS = ["test", "Test", "Mock", "mock"]


def get_file_type(filename):
    if filename.endswith(".java"):   return "JAVA"
    elif filename.endswith(".xml"): return "LAYOUT XML"
    return "CODE"


def should_include(file, root):
    if file.endswith(".java"):
        if any(k in file for k in EXCLUDED_KEYWORDS): return False
        if "src" + os.sep + "main" not in root:       return False
        return True
    if file.endswith(".xml"):
        return os.path.basename(root).startswith("layout")
    return False


def generate_doc(app_dir, output_file, log_cb, done_cb):
    seen = set()
    doc  = Document()
    doc.add_heading("Project Code Documentation", 0)
    doc.add_paragraph(app_dir)
    doc.add_page_break()
    count = 0

    for root, dirs, files in os.walk(app_dir):
        dirs[:] = [d for d in dirs if d not in EXCLUDED_DIRS]
        for file in sorted(files):
            if not should_include(file, root): continue
            if file in seen:
                log_cb(f"⚠  דילוג כפילות: {file}")
                continue
            seen.add(file)
            file_path = os.path.join(root, file)
            file_type = get_file_type(file)
            log_cb(f"  {file}  [{file_type}]")
            h = doc.add_heading(level=1)
            h.add_run(f"📄 {file}  [{file_type}]").bold = True
            doc.add_paragraph("─" * 40)
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        p = doc.add_paragraph()
                        r = p.add_run(line.rstrip())
                        r.font.name = "Courier New"
                        r.font.size = Pt(9)
            except Exception as e:
                doc.add_paragraph(f"שגיאה: {e}")
            doc.add_paragraph()
            count += 1

    doc.save(output_file)
    done_cb(count, output_file)


# ─────────────────────────────
# GUI – סגנון Auto Clicker פשוט
# ─────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Project Book Generator")
        self.geometry("500x460")
        self.resizable(False, False)
        self.configure(bg="#f0f0f0")
        self._build()

    def _build(self):
        # ── כותרת ──
        title_bar = tk.Frame(self, bg="#2d6cdf", height=42)
        title_bar.pack(fill="x")
        tk.Label(title_bar, text="Project Book Generator",
                 bg="#2d6cdf", fg="white",
                 font=("Segoe UI", 12, "bold")).pack(side="left", padx=14, pady=8)

        pad = dict(padx=16, pady=6)

        # ── נתיב APP ──
        tk.Label(self, text="App Folder:", bg="#f0f0f0",
                 font=("Segoe UI", 9, "bold"), anchor="w").pack(fill="x", **pad)

        row1 = tk.Frame(self, bg="#f0f0f0")
        row1.pack(fill="x", padx=16, pady=(0, 8))
        self.dir_var = tk.StringVar()
        tk.Entry(row1, textvariable=self.dir_var, font=("Segoe UI", 9),
                 relief="sunken", bd=2).pack(side="left", fill="x", expand=True)
        tk.Button(row1, text="Browse...", font=("Segoe UI", 9),
                  bg="#e1e1e1", relief="raised", bd=2, width=9,
                  cursor="hand2", command=self._browse_dir).pack(side="left", padx=(6,0))

        # ── קובץ פלט ──
        tk.Label(self, text="Output File:", bg="#f0f0f0",
                 font=("Segoe UI", 9, "bold"), anchor="w").pack(fill="x", **pad)

        row2 = tk.Frame(self, bg="#f0f0f0")
        row2.pack(fill="x", padx=16, pady=(0, 12))
        self.out_var = tk.StringVar(value="ProjectBook.docx")
        tk.Entry(row2, textvariable=self.out_var, font=("Segoe UI", 9),
                 relief="sunken", bd=2).pack(side="left", fill="x", expand=True)
        tk.Button(row2, text="Save As...", font=("Segoe UI", 9),
                  bg="#e1e1e1", relief="raised", bd=2, width=9,
                  cursor="hand2", command=self._browse_out).pack(side="left", padx=(6,0))

        # ── קו הפרדה ──
        ttk.Separator(self, orient="horizontal").pack(fill="x", padx=16, pady=4)

        # ── כפתור הפעלה ──
        self.run_btn = tk.Button(self, text="▶  Generate Document",
                                 font=("Segoe UI", 10, "bold"),
                                 bg="#2d6cdf", fg="white",
                                 activebackground="#1a52b8", activeforeground="white",
                                 relief="raised", bd=2, height=2,
                                 cursor="hand2", command=self._run)
        self.run_btn.pack(fill="x", padx=16, pady=8)

        # ── progress bar ──
        self.progress = ttk.Progressbar(self, mode="indeterminate", length=468)
        self.progress.pack(padx=16, pady=(0, 8))

        # ── לוג ──
        tk.Label(self, text="Log:", bg="#f0f0f0",
                 font=("Segoe UI", 9, "bold"), anchor="w").pack(fill="x", padx=16)

        log_frame = tk.Frame(self, bg="#f0f0f0")
        log_frame.pack(fill="both", expand=True, padx=16, pady=(2, 14))

        self.log = tk.Text(log_frame, font=("Consolas", 8),
                           bg="white", fg="#222", relief="sunken", bd=2,
                           state="disabled", height=10)
        sb = tk.Scrollbar(log_frame, command=self.log.yview)
        self.log.configure(yscrollcommand=sb.set)
        self.log.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # ── סטטוס בר ──
        self.status_var = tk.StringVar(value="Ready")
        tk.Label(self, textvariable=self.status_var,
                 bg="#d4d4d4", fg="#333", anchor="w",
                 font=("Segoe UI", 8), relief="sunken", bd=1).pack(
                     fill="x", side="bottom", ipady=2, padx=0)

    # ── עזר ──
    def _browse_dir(self):
        p = filedialog.askdirectory(title="Select app folder")
        if p: self.dir_var.set(p)

    def _browse_out(self):
        p = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile="ProjectBook.docx")
        if p: self.out_var.set(p)

    def _log(self, msg):
        self.log.configure(state="normal")
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.log.configure(state="disabled")
        self.status_var.set(msg[:80])

    def _run(self):
        app_dir = self.dir_var.get().strip()
        out     = self.out_var.get().strip()
        if not app_dir or not os.path.isdir(app_dir):
            messagebox.showerror("Error", "Please select a valid app folder.")
            return
        if not out:
            messagebox.showerror("Error", "Please enter an output filename.")
            return

        self.log.configure(state="normal")
        self.log.delete("1.0", "end")
        self.log.configure(state="disabled")

        self.run_btn.configure(state="disabled", text="⏳  Working...")
        self.progress.start(10)
        self._log(f"Scanning: {app_dir}")

        threading.Thread(
            target=generate_doc,
            args=(app_dir, out, self._log,
                  lambda c, p: self.after(0, self._done, c, p)),
            daemon=True
        ).start()

    def _done(self, count, output_file):
        self.progress.stop()
        self.run_btn.configure(state="normal", text="▶  Generate Document")
        self._log(f"\nDone!  {count} files → {output_file}")
        self.status_var.set(f"Done — {count} files written to {output_file}")
        messagebox.showinfo("Done!", f"Created: {output_file}\n{count} files documented.")


if __name__ == "__main__":
    App().mainloop()