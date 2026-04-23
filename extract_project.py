"""
extract_project.py
------------------
מייצר קובץ Word מתועד מתוך פרויקט Android.
סורק:
  - כל קבצי *.java
  - קבצי *.xml שנמצאים בתוך res/layout בלבד

שימוש:
    python extract_project.py <PATH_TO_APP_FOLDER>

דוגמא:
    python extract_project.py C:/MyProject/app
"""

import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────
# הגדרות
# ──────────────────────────────────────────
OUTPUT_FILE = "ProjectBook.docx"


def parse_args():
    parser = argparse.ArgumentParser(
        description="מייצר ספר קוד Word מתוך פרויקט Android"
    )
    parser.add_argument(
        "app_path",
        nargs="?",
        default=None,
        help="נתיב לתיקיית ה-app של הפרויקט",
    )
    return parser.parse_args()


def get_app_path(arg_path: str | None) -> Path:
    """מקבל נתיב מ-argument או שואל את המשתמש."""
    if arg_path:
        p = Path(arg_path)
    else:
        raw = input("📁 הכנס נתיב לתיקיית ה-app של הפרויקט: ").strip().strip('"').strip("'")
        p = Path(raw)

    if not p.exists():
        print(f"❌ שגיאה: הנתיב '{p}' לא קיים.")
        sys.exit(1)
    if not p.is_dir():
        print(f"❌ שגיאה: '{p}' אינו תיקייה.")
        sys.exit(1)
    return p.resolve()


def is_layout_xml(file_path: Path, app_root: Path) -> bool:
    """בודק אם קובץ XML נמצא בתוך res/layout (בכל רמה עומק)."""
    try:
        relative = file_path.relative_to(app_root)
        parts = relative.parts
        for i, part in enumerate(parts[:-1]):  # לא כולל שם הקובץ עצמו
            if part == "res" and i + 1 < len(parts) - 1:
                next_part = parts[i + 1]
                # res/layout  או  res/layout-land  וכו'
                if next_part == "layout" or next_part.startswith("layout"):
                    return True
    except ValueError:
        pass
    return False


def collect_files(app_root: Path) -> list[Path]:
    """
    מאסף את כל קבצי Java + XML של layout,
    ממיין לפי סוג ואחר-כך לפי נתיב.
    """
    java_files = []
    layout_files = []

    for root, dirs, files in os.walk(app_root):
        # דלג על תיקיות build
        dirs[:] = [d for d in dirs if d not in ("build", ".gradle", ".idea", "captures")]
        for filename in files:
            fp = Path(root) / filename
            if filename.endswith(".java"):
                java_files.append(fp)
            elif filename.endswith(".xml") and is_layout_xml(fp, app_root):
                layout_files.append(fp)

    java_files.sort()
    layout_files.sort()
    return java_files + layout_files


def add_horizontal_rule(doc: Document):
    """מוסיף קו הפרדה אופקי."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_file_section(doc: Document, file_path: Path, app_root: Path):
    """מוסיף פרק לקובץ יחיד: כותרת + מטה-דאטה + קוד."""
    relative = file_path.relative_to(app_root)
    suffix = file_path.suffix.lower()

    file_type = "JAVA" if suffix == ".java" else "XML / LAYOUT"

    # ── כותרת קובץ (Heading 1) ─────────────────────
    heading = doc.add_heading(level=1)
    heading.clear()
    run = heading.add_run(f"📄  {file_path.name}   [{file_type}]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # ── נתיב יחסי ─────────────────────────────────
    info_para = doc.add_paragraph()
    label_run = info_para.add_run("📂 נתיב:  ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    path_run = info_para.add_run(str(relative))
    path_run.font.name = "Courier New"
    path_run.font.size = Pt(9)
    path_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    add_horizontal_rule(doc)

    # ── קוד ────────────────────────────────────────
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()

        if not lines:
            p = doc.add_paragraph()
            p.add_run("(קובץ ריק)").italic = True
        else:
            for line in lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.rstrip("\n\r"))
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    except Exception as e:
        err_para = doc.add_paragraph()
        err_para.add_run(f"⚠️  שגיאה בקריאת הקובץ: {e}").font.color.rgb = RGBColor(0xCC, 0, 0)

    # ── רווח בין קבצים ─────────────────────────────
    doc.add_paragraph()


def build_document(app_root: Path, files: list[Path]) -> Document:
    doc = Document()

    # ── כותרת ראשית ────────────────────────────────
    title = doc.add_heading("Project Code Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"📁 App Path:  {app_root}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── סטטיסטיקה ──────────────────────────────────
    java_count = sum(1 for f in files if f.suffix == ".java")
    xml_count = sum(1 for f in files if f.suffix == ".xml")
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_run = stats.add_run(
        f"📊 {len(files)} קבצים  ·  {java_count} Java  ·  {xml_count} Layout XML"
    )
    stats_run.font.size = Pt(10)
    stats_run.bold = True

    doc.add_page_break()

    # ── קבצים ──────────────────────────────────────
    for fp in files:
        add_file_section(doc, fp, app_root)

    return doc


def main():
    args = parse_args()
    app_root = get_app_path(args.app_path)

    print(f"\n🔍 סורק: {app_root}")
    files = collect_files(app_root)

    if not files:
        print("⚠️  לא נמצאו קבצי Java או Layout XML בנתיב שצוין.")
        sys.exit(0)

    java_count = sum(1 for f in files if f.suffix == ".java")
    xml_count = sum(1 for f in files if f.suffix == ".xml")
    print(f"✅ נמצאו: {java_count} קבצי Java, {xml_count} קבצי Layout XML")
    print("📝 מייצר מסמך Word...")

    doc = build_document(app_root, files)
    output_path = Path(OUTPUT_FILE)
    doc.save(output_path)

    print(f"\n🎉 הושלם! נוצר הקובץ: {output_path.resolve()}")


if __name__ == "__main__":
    main()
